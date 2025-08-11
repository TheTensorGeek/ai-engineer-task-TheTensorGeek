from docx import Document

KEYWORD_MAP = {
    'articles of association': ['articles of association', 'aoa'],
    'memorandum of association': ['memorandum of association', 'moa', 'memorandum'],
    'board resolution': ['board resolution', 'resolution of the board'],
    'ubo declaration': ['ubo declaration', 'ultimate beneficial owner', 'ubo'],
    'register of members and directors': ['register of members', 'register of directors', 'register']
}

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(paragraphs)

def classify_document(text: str) -> str:
    t = text.lower()
    for dtype, kws in KEYWORD_MAP.items():
        for kw in kws:
            if kw in t[:2000]:
                return dtype
    if 'articles' in t[:1000]:
        return 'articles of association'
    if 'memorandum' in t[:1000]:
        return 'memorandum of association'
    return 'unknown'

def add_inline_comment(input_path: str, output_path: str, comment_annotations):
    doc = Document(input_path)
    for ann in comment_annotations:
        snippet = ann.get('snippet') or ''
        message = ann.get('message') or ''
        for para in doc.paragraphs:
            if snippet and snippet in para.text:
                para.text = para.text.replace(snippet, f"{snippet} [COMMENT: {message}]")
    doc.save(output_path)
