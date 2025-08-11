import os, tempfile
from doc_parser import extract_text_from_docx, classify_document, add_inline_comment
from rag_pipeline import load_index, retrieve, analyze_with_llm
from checklist import match_uploaded_to_required
from utils import write_summary, timestamp

import gradio as gr
import json
from datetime import datetime
from docx import Document

# ======== Safe Save Function ========
def save_outputs(ai_review_text, ai_json_summary):
    # Timestamp for unique file names
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Paths
    output_docx_path = f"/content/reviewed_{timestamp}.docx"
    output_json_path = f"/content/review_summary_{timestamp}.json"

    # ---------- 1) Save reviewed.docx ----------
    try:
        doc = Document()
        doc.add_heading("AI Review", level=1)
        if isinstance(ai_review_text, str):
            doc.add_paragraph(ai_review_text)
        else:
            doc.add_paragraph(str(ai_review_text))  # Fallback if not a string
        doc.save(output_docx_path)
        print(f"✅ Saved reviewed.docx -> {output_docx_path}")
    except Exception as e:
        print(f"⚠ Error saving reviewed.docx: {e}")

    # ---------- 2) Save JSON summary ----------
    try:
        if isinstance(ai_json_summary, str):
            try:
                parsed_json = json.loads(ai_json_summary)
            except json.JSONDecodeError:
                print("⚠ JSON parsing failed. Saving as raw text instead.")
                parsed_json = {"raw_output": ai_json_summary}
        elif isinstance(ai_json_summary, dict):
            parsed_json = ai_json_summary
        else:
            parsed_json = {"raw_output": str(ai_json_summary)}

        with open(output_json_path, "w", encoding="utf-8") as f:
            json.dump(parsed_json, f, indent=2, ensure_ascii=False)
        print(f"✅ Saved review_summary.json -> {output_json_path}")

    except Exception as e:
        print(f"⚠ Error saving JSON summary: {e}")

    return output_docx_path, output_json_path

# ======== Your Main Review Function ========
def review_docx(file):
    # --- 1) Load the uploaded file ---
    file_path = file.name

    # --- 2) Run your AI logic here ---
    # Replace this with your actual processing logic
    ai_review_text = "This is a sample reviewed text from AI."
    ai_json_summary = {"summary": "This is a sample JSON summary from AI."}

    # --- 3) Save files safely ---
    docx_path, json_path = save_outputs(ai_review_text, ai_json_summary)

    return docx_path, json_path

# ======== Gradio UI ========
with gr.Blocks() as demo:
    gr.Markdown("## 📄 ADGM AOA Review Tool")
    file_input = gr.File(label="Upload DOCX File", file_types=[".docx"])
    docx_output = gr.File(label="Reviewed DOCX")
    json_output = gr.File(label="Review Summary JSON")

    file_input.change(fn=review_docx, inputs=file_input, outputs=[docx_output, json_output])

# ======== Launch Gradio ========
if __name__ == "__main__":
    demo.launch(server_name="0.0.0.0", server_port=7860, allowed_paths=["/content"])
